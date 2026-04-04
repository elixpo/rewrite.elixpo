"""Document parsing — PDF and DOCX extraction to structured Document."""

import logging
import re
from pathlib import Path

from app.document.structure import Document, Section, Paragraph

logger = logging.getLogger(__name__)


def parse_pdf(path: str) -> Document:
    """Extract text from PDF preserving basic structure."""
    from PyPDF2 import PdfReader

    reader = PdfReader(path)
    pages_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages_text.append(text.strip())

    full_text = "\n\n".join(pages_text)
    return _text_to_document(full_text, title=Path(path).stem, source_path=path)


def parse_docx(path: str) -> Document:
    """Extract text from DOCX preserving headings and structure."""
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    sections = []
    current_heading = ""
    current_paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect headings by style
        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            # Flush current section
            if current_paragraphs:
                sections.append(Section(
                    heading=current_heading,
                    paragraphs=[Paragraph(text=p) for p in current_paragraphs],
                ))
                current_paragraphs = []
            current_heading = text
        else:
            current_paragraphs.append(text)

    # Flush last section
    if current_paragraphs:
        sections.append(Section(
            heading=current_heading,
            paragraphs=[Paragraph(text=p) for p in current_paragraphs],
        ))

    if not sections:
        # Fallback: treat as single section
        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        return _text_to_document(full_text, title=Path(path).stem, source_path=path)

    return Document(title=Path(path).stem, sections=sections, source_path=path)


def parse_tex(path: str) -> Document:
    """Extract text from LaTeX (.tex) preserving section structure.

    Strips LaTeX commands, preserves section headings and paragraph text.
    Handles math environments, citations, and common LaTeX patterns.
    """
    with open(path, "r", encoding="utf-8") as f:
        raw = f.read()

    # Strip comments (lines starting with % and inline comments)
    raw = re.sub(r"(?<!\\)%.*$", "", raw, flags=re.MULTILINE)

    # Extract title if present
    title_match = re.search(r"\\title\{([^}]*)\}", raw)
    title = _strip_tex_commands(title_match.group(1)) if title_match else Path(path).stem

    # Extract body between \begin{document} and \end{document}
    body_match = re.search(
        r"\\begin\{document\}(.*?)\\end\{document\}", raw, re.DOTALL
    )
    body = body_match.group(1) if body_match else raw

    # Remove common preamble commands that leak through
    body = re.sub(r"\\maketitle", "", body)
    body = re.sub(r"\\tableofcontents", "", body)
    body = re.sub(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", r"\1", body, flags=re.DOTALL)

    # Parse sections
    # Match \section{}, \subsection{}, \subsubsection{}, and starred variants
    section_pattern = re.compile(
        r"\\(section|subsection|subsubsection)\*?\{([^}]*)\}", re.DOTALL
    )

    sections = []
    last_end = 0
    current_heading = ""

    for match in section_pattern.finditer(body):
        # Flush text before this heading
        pre_text = body[last_end:match.start()]
        pre_text = _clean_tex_body(pre_text)
        if pre_text.strip():
            paragraphs = _split_tex_paragraphs(pre_text)
            if paragraphs:
                sections.append(Section(heading=current_heading, paragraphs=paragraphs))

        current_heading = _strip_tex_commands(match.group(2))
        last_end = match.end()

    # Flush remaining text after last section heading
    remaining = body[last_end:]
    remaining = _clean_tex_body(remaining)
    if remaining.strip():
        paragraphs = _split_tex_paragraphs(remaining)
        if paragraphs:
            sections.append(Section(heading=current_heading, paragraphs=paragraphs))

    if not sections:
        # No sections found — treat whole body as one section
        cleaned = _clean_tex_body(body)
        paragraphs = _split_tex_paragraphs(cleaned)
        sections = [Section(heading="", paragraphs=paragraphs or [Paragraph(text=cleaned.strip())])]

    return Document(title=title, sections=sections, source_path=path)


def _clean_tex_body(text: str) -> str:
    """Strip LaTeX commands from body text, preserving readable content."""
    result = text

    # Remove figure/table/listing environments entirely
    for env in ["figure", "table", "listing", "tikzpicture", "algorithm"]:
        result = re.sub(
            rf"\\begin\{{{env}\}}.*?\\end\{{{env}\}}", "", result, flags=re.DOTALL
        )

    # Replace display math with placeholder
    result = re.sub(r"\\\[.*?\\\]", " [equation] ", result, flags=re.DOTALL)
    result = re.sub(r"\\begin\{equation\*?\}.*?\\end\{equation\*?\}", " [equation] ", result, flags=re.DOTALL)
    result = re.sub(r"\\begin\{align\*?\}.*?\\end\{align\*?\}", " [equation] ", result, flags=re.DOTALL)
    result = re.sub(r"\\begin\{gather\*?\}.*?\\end\{gather\*?\}", " [equation] ", result, flags=re.DOTALL)

    # Replace inline math with content or placeholder
    result = re.sub(r"\$([^$]+)\$", r" \1 ", result)

    # Handle common text commands — extract their content
    for cmd in ["textbf", "textit", "emph", "underline", "texttt", "text", "mathrm"]:
        result = re.sub(rf"\\{cmd}\{{([^}}]*)\}}", r"\1", result)

    # Handle \cite{...} → [citation]
    result = re.sub(r"\\cite[tp]?\*?\{[^}]*\}", "[citation]", result)

    # Handle \ref{...} and \label{...}
    result = re.sub(r"\\(?:ref|eqref|autoref|cref)\{[^}]*\}", "[ref]", result)
    result = re.sub(r"\\label\{[^}]*\}", "", result)

    # Handle \footnote{...} — keep the text
    result = re.sub(r"\\footnote\{([^}]*)\}", r" (\1)", result)

    # Handle itemize/enumerate — convert items to plain text
    result = re.sub(r"\\begin\{(?:itemize|enumerate)\}", "", result)
    result = re.sub(r"\\end\{(?:itemize|enumerate)\}", "", result)
    result = re.sub(r"\\item\s*", "- ", result)

    # Strip remaining commands with no arguments
    result = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^}]*)\}", r"\1", result)
    result = re.sub(r"\\[a-zA-Z]+\*?", " ", result)

    # Clean up braces, multiple spaces, etc.
    result = result.replace("{", "").replace("}", "")
    result = re.sub(r"[ \t]+", " ", result)
    result = re.sub(r"\n{3,}", "\n\n", result)

    return result.strip()


def _strip_tex_commands(text: str) -> str:
    """Strip TeX commands from a short string (heading, title)."""
    result = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", text)
    result = re.sub(r"\\[a-zA-Z]+", "", result)
    result = result.replace("{", "").replace("}", "")
    return result.strip()


def _split_tex_paragraphs(text: str) -> list:
    """Split cleaned TeX text into Paragraph objects on double newlines."""
    raw = [p.strip() for p in text.split("\n\n") if p.strip()]
    # Filter out very short fragments (likely leftover commands)
    return [Paragraph(text=p) for p in raw if len(p) > 20]


def parse_file(path: str) -> Document:
    """Auto-detect file type and parse. Supports .pdf, .docx, .tex, .txt."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(path)
    elif ext == ".docx":
        return parse_docx(path)
    elif ext == ".tex":
        return parse_tex(path)
    elif ext in (".txt", ".md"):
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        return _text_to_document(text, title=Path(path).stem, source_path=path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .pdf, .docx, .tex, or .txt")


def _text_to_document(text: str, title: str = "", source_path: str = None) -> Document:
    """Convert raw text into a structured Document with section detection."""
    lines = text.split("\n")
    sections = []
    current_heading = ""
    current_lines = []

    # Simple heading detection: short lines (< 80 chars) that are all caps or
    # followed by blank lines, or lines starting with common section patterns
    heading_pattern = re.compile(
        r"^(?:\d+\.?\s+)?(?:abstract|introduction|background|methods?|results?|"
        r"discussion|conclusion|references|acknowledgments?|appendix)",
        re.IGNORECASE,
    )

    for line in lines:
        stripped = line.strip()

        is_heading = (
            stripped
            and len(stripped) < 80
            and (
                stripped.isupper()
                or heading_pattern.match(stripped)
                or (stripped.startswith("#") and len(stripped) < 100)
            )
        )

        if is_heading:
            # Flush current section
            paragraph_text = "\n".join(current_lines).strip()
            if paragraph_text:
                paragraphs = [
                    Paragraph(text=p.strip())
                    for p in paragraph_text.split("\n\n")
                    if p.strip()
                ]
                sections.append(Section(heading=current_heading, paragraphs=paragraphs))
                current_lines = []
            current_heading = stripped.lstrip("# ")
        else:
            current_lines.append(line)

    # Flush last
    paragraph_text = "\n".join(current_lines).strip()
    if paragraph_text:
        paragraphs = [
            Paragraph(text=p.strip())
            for p in paragraph_text.split("\n\n")
            if p.strip()
        ]
        sections.append(Section(heading=current_heading, paragraphs=paragraphs))

    if not sections:
        sections = [Section(heading="", paragraphs=[Paragraph(text=text.strip())])]

    return Document(title=title, sections=sections, source_path=source_path)
